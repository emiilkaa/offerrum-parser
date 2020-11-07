import getpass
import os
import re
import shutil
import time

import docx
import mammoth
import requests
from bs4 import BeautifulSoup
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from selenium import webdriver
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait


def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


def page_source(driver):
    return driver.execute_script('return document.documentElement.outerHTML;')


driver = webdriver.Firefox()
driver.set_page_load_timeout(60)

driver.get('https://offerrum.com/')
email = input('Email: ')
attempts = 1
while not re.match(r'\w+@\w+\.\w+', email):
    email = input('Invalid format. Enter your email again: ')
    attempts += 1
    if attempts == 5:
        break
if not re.match(r'\w+@\w+\.\w+', email):
    driver.close()
    raise Exception('Invalid email format')
password = getpass.getpass()
driver.find_element_by_name('email').send_keys(email)
driver.find_element_by_name('password').send_keys(password)
try:
    driver.find_element_by_class_name('popup-close').click()
except:
    pass
driver.find_element_by_name('password').send_keys(Keys.RETURN)
time.sleep(3)

if 'Email или пароль ошибочные' in page_source(driver):
    driver.close()
    raise Exception('Email или пароль ошибочные')
document = docx.Document()

with open('links.txt', 'r', encoding='utf-8') as file:
    links = [line.strip() for line in file]

for i in range(len(links)):
    links[i] = re.search(r'\d{2,}', links[i]).group(0)
    print('ID оффера: ' + links[i] + '\n')
    url = 'https://my.offerrum.com/offers/' + links[i]
    href = '/tickets/create?theme=pixelRequest&offer=' + links[i]
    ok = 1
    for k in range(5):
        try:
            driver.get(url)
            element = WebDriverWait(driver, 15).until(
                EC.presence_of_element_located((By.XPATH, '//a[@href="' + href + '"]'))
            )
        except TimeoutException as e:
            print('Failed to connect')
            if k == 4:
                print()
                ok = 0
        else:
            break
    if not ok:
        continue
    data = page_source(driver)
    soup = BeautifulSoup(data, 'lxml')
    title = soup.find('h3', class_='offer-title').get_text()
    title_word = document.add_paragraph(str(i + 1) + '. Оффер: ')
    title_word.add_run(title).bold = True
    reward = []
    for t in soup.find_all('span', class_='pixel-reward'):
        reward.append(t)
    if '₽' in str(reward):
        currency = ' рублей'
    elif '$' in str(reward):
        currency = '$'
    elif '€' in str(reward):
        currency = '€'
    for t in range(len(reward)):
        if '.5' in str(reward[t]):
            reward[t] = float(re.search(r'\d+\.5', str(reward[t])).group(0))
        else:
            reward[t] = int(re.search(r'\d+', str(reward[t])).group(0))
    reward = list(set(reward))
    reward_word = document.add_paragraph('Отчисления: ')
    if len(reward) == 1:
        reward_word.add_run(str(reward[0]) + currency)
    else:
        reward_word.add_run('от ' + str(min(reward)) + ' до ' + str(max(reward)) + currency)
    ul = soup.find(class_='collapsible-list')
    d = {li.span.string: li.span.find_next('span').string for li in ul.find_all('li')}
    geos = list(d.keys())
    prices = list(set(d.values()))
    if len(prices) == 1:
        price = prices[0]
    elif len(prices) == 2:
        price = prices[0] + ' или ' + prices[1]
    else:
        prices[-1] = 'или ' + prices[-1]
        price = ', '.join(prices)
        price = price.replace(', или', ' или')
    price_word = document.add_paragraph('Цена: ' + price)
    geo = ', '.join(geos)
    geo_word = document.add_paragraph('ГЕО: ' + geo)
    cpa = document.add_paragraph('CPA-сеть: ')
    add_hyperlink(cpa, 'OFFERRUM', 'https://offerrum.com/')
    name = links[i] + '.png'
    for k in range(5):
        try:
            picture = requests.get(
                'https://static.offerrum.com/networks/ad1/img/offers_covers/' + links[i] + '_r_original.png',
                stream=True, timeout=30)
            if picture.status_code == 200:
                with open(name, 'wb') as f:
                    picture.raw.decode_content = True
                    shutil.copyfileobj(picture.raw, f)
            document.add_picture(name)
            os.remove(name)
        except:
            print('Failed to download')
            if k == 4:
                print()
        else:
            break

driver.close()
filename = input('Please enter a filename: ')
if len(filename) >= 5 and filename[-5:] in ('.docx', '.html'):
    filename = filename[:-5]
if not filename:
    filename = 'OFFERRUM'
document.save(filename + '.docx')
with open(filename + '.docx', 'rb') as f, open(filename + '.html', 'wb') as b:
    page = mammoth.convert_to_html(f)
    b.write(page.value.encode('utf-8'))
